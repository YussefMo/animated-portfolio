"""
CV Generator — creates tailored .docx CVs and converts to PDF.

Usage:
  python scripts/generate_cv.py                  # base CV → youssef_cv.docx
  python scripts/generate_cv.py --variant react  # React-focused variant
  python scripts/generate_cv.py --pdf            # also output PDF
  python scripts/generate_cv.py --job "Frontend Engineer at Vercel"
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import argparse, os, json

# ─── DATA ───────────────────────────────────────────────

NAME = "Youssef Mohammed Ali"
TITLE = "Frontend Engineer"
EMAIL = "youssefmohammed12090@gmail.com"
PHONE = "+20 1062479462"
LOCATION = "Menoufia, Egypt"
PORTFOLIO = "https://www.ymdev.me"
LINKEDIN = "https://linkedin.com/in/ymdev"
GITHUB = "https://github.com/YussefMo"

SUMMARY_BASE = (
    "Frontend Engineer with production experience building scalable SaaS web "
    "and mobile applications using React, Next.js, TypeScript, and React Native. "
    "Experienced in modern frontend architecture, performance optimization, "
    "reusable design systems, AI-powered user experiences, and delivering "
    "measurable business impact."
)

EXPERIENCES = [
    {
        "company": "Diagnosit",
        "role": "Frontend Engineer",
        "period": "Jun 2025 — Present",
        "location": "Remote, UAE",
        "bullets": [
            "Re-architected a large Next.js SaaS platform into modular architecture, accelerating feature delivery by ~50% and improving long-term maintainability.",
            "Led a complete UI/UX redesign that reduced rendering issues by ~50% and established consistent, accessible design patterns across the product.",
            "Built the Live Patient Link feature, reducing initial load time by 45% and bundle size by 70%.",
            "Optimized landing page performance, cutting load times by 60% through render optimization, asset pipeline improvements, and strategic code splitting.",
            "Architected the React Native (Expo) mobile application with full feature parity to the web platform — authentication, real-time dashboards, push notifications (FCM).",
            "Designed and shipped a reusable design system adopted across the product, reducing new feature development time by 45%.",
            "Integrated and maintained 30+ REST API endpoints with TanStack Query and Axios, implementing type-safe data fetching, optimistic updates, and cache invalidation.",
            "Implemented full internationalization (i18n) across web and mobile with locale-aware routing, translated UI, and RTL layout support.",
        ]
    },
    {
        "company": "Web Masters",
        "role": "Frontend Developer Intern",
        "period": "Apr 2025 — Jun 2025",
        "location": "Remote, Egypt",
        "bullets": [
            "Built responsive React and Next.js interfaces from Figma designs within an Agile team.",
            "Refactored legacy components into clean, maintainable, production-ready code.",
            "Collaborated through Git workflows, code reviews, and version control best practices.",
        ]
    },
]

EDUCATION = {
    "degree": "Bachelor of Business Information Systems (BIS)",
    "school": "El-Gezera Institute",
    "period": "Oct 2021 — Jun 2025",
}

SKILLS = [
    "JavaScript (ES6), TypeScript, HTML5, CSS3",
    "React, Next.js, React Native, Expo",
    "TanStack Query, Axios, REST APIs",
    "Tailwind CSS, Design Systems, Component Architecture",
    "Firebase, Supabase",
    "Git, GitHub, Figma, Vercel",
]

# ─── VARIANTS ───────────────────────────────────────────

VARIANTS = {
    "react": {
        "title": "React / Next.js Frontend Engineer",
        "summary_extra": (
            "Deep expertise in React Server Components, App Router architecture, "
            "and Next.js performance optimization. Built and shipped production "
            "SaaS applications with measurable performance gains."
        ),
    },
    "mobile": {
        "title": "React Native / Mobile Frontend Engineer",
        "summary_extra": (
            "Specialized in cross-platform mobile development with React Native and Expo. "
            "Architected mobile apps with full feature parity to web platforms — "
            "authentication, real-time dashboards, push notifications, and i18n."
        ),
    },
    "ai": {
        "title": "AI Frontend Engineer",
        "summary_extra": (
            "Experienced in building AI-powered frontend experiences: LLM integrations, "
            "conversational agent interfaces, tool-augmented chatbots, and real-time "
            "streaming response UIs."
        ),
    },
}

# ─── BUILD ──────────────────────────────────────────────

def build_cv(variant=None, job_description=None, output_path="youssef_cv.docx"):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    # ── Header ──
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(NAME)
    run.bold = True
    run.size = Pt(20)
    run.font.color.rgb = RGBColor(0x0d, 0x0d, 0x0d)

    variant_title = VARIANTS.get(variant, {}).get("title", TITLE) if variant else TITLE
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(variant_title)
    run.size = Pt(12)
    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(2)
    contact_info = f"{LOCATION}  |  {EMAIL}  |  {PHONE}"
    run = contact_p.add_run(contact_info)
    run.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    links_p = doc.add_paragraph()
    links_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links_p.paragraph_format.space_after = Pt(10)
    links_text = f"LinkedIn: {LINKEDIN}  |  GitHub: {GITHUB}  |  Portfolio: {PORTFOLIO}"
    run = links_p.add_run(links_text)
    run.size = Pt(9)
    run.font.color.rgb = RGBColor(0x4b, 0x4d, 0x83)

    add_section_heading(doc, "")

    # ── Summary ──
    add_section_heading(doc, "Professional Summary")
    summary = SUMMARY_BASE
    if variant and variant in VARIANTS:
        extra = VARIANTS[variant].get("summary_extra", "")
        if extra:
            summary = SUMMARY_BASE + " " + extra
    if job_description:
        summary = tailor_summary(summary, job_description)
    p = doc.add_paragraph(summary)
    p.paragraph_format.space_after = Pt(8)

    # ── Experience ──
    add_section_heading(doc, "Experience")
    for exp in EXPERIENCES:
        header = doc.add_paragraph()
        header.paragraph_format.space_after = Pt(0)
        header.paragraph_format.space_before = Pt(6)
        run = header.add_run(f"{exp['company']} — {exp['role']}")
        run.bold = True
        run.size = Pt(10.5)

        meta = doc.add_paragraph()
        meta.paragraph_format.space_after = Pt(2)
        run = meta.add_run(f"{exp['location']}  |  {exp['period']}")
        run.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        for bullet in exp["bullets"]:
            bp = doc.add_paragraph(style="List Bullet")
            bp.clear()
            run = bp.add_run(bullet)
            run.size = Pt(10)
            bp.paragraph_format.space_after = Pt(2)

    # ── Education ──
    add_section_heading(doc, "Education")
    edu = doc.add_paragraph()
    edu.paragraph_format.space_after = Pt(2)
    run = edu.add_run(f"{EDUCATION['degree']} — {EDUCATION['school']}")
    run.bold = True
    run.size = Pt(10.5)

    edu_meta = doc.add_paragraph()
    edu_meta.paragraph_format.space_after = Pt(8)
    run = edu_meta.add_run(EDUCATION["period"])
    run.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── Skills ──
    add_section_heading(doc, "Technical Skills")
    for skill in SKILLS:
        sp = doc.add_paragraph(skill)
        sp.paragraph_format.space_after = Pt(1)
        sp.style.font.size = Pt(10)

    doc.save(output_path)
    print(f"Saved: {os.path.abspath(output_path)}")
    return output_path


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    if text:
        run = p.add_run(text.upper())
        run.bold = True
        run.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    else:
        p.paragraph_format.space_before = Pt(0)
    return p


def tailor_summary(summary, job_description):
    keywords = job_description.lower()
    additions = []
    if "react native" in keywords or "mobile" in keywords:
        additions.append("Proven experience architecting cross-platform mobile applications with React Native and Expo.")
    if "ai" in keywords or "llm" in keywords or "machine learning" in keywords:
        additions.append("Hands-on experience integrating AI/LLM features into production frontend applications.")
    if "design system" in keywords:
        additions.append("Shipped reusable design systems that accelerated team velocity by 45%.")
    if "performance" in keywords:
        additions.append("Track record of measurable performance optimization — 60% faster page loads, 70% smaller bundles.")
    if additions:
        return summary + " " + " ".join(additions)
    return summary


# ─── CLI ────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate tailored CV")
    parser.add_argument("--variant", choices=list(VARIANTS.keys()), help="CV variant")
    parser.add_argument("--job", type=str, help="Job description for tailoring")
    parser.add_argument("--output", type=str, default=None, help="Output path")
    parser.add_argument("--pdf", action="store_true", help="Also convert to PDF")

    args = parser.parse_args()

    output = args.output or (
        f"youssef_cv_{args.variant}.docx" if args.variant else "youssef_cv.docx"
    )

    path = build_cv(variant=args.variant, job_description=args.job, output_path=output)

    if args.pdf:
        import subprocess
        pdf_path = output.replace(".docx", ".pdf")
        script = os.path.join(os.path.dirname(__file__), "docx2pdf.js")
        subprocess.run(["node", script, output, pdf_path], check=True)
